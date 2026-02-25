# -*- coding: utf-8 -*-
"""
extract_fiches.py â€” Extraction de fiches Word structurÃ©es
==========================================================
Lit un fichier .docx structurÃ© avec des balises de la forme :
  Nom commun: Ortie
  Type: plante brute
  PropriÃ©tÃ©s: DiurÃ©tique, reminÃ©ralisant...

Retourne un objet Plante du bon type, prÃªt Ã  Ãªtre sauvegardÃ© via database.py.

Structure attendue des fiches Word :
  - Chaque champ sur sa propre ligne : "Label: valeur"
  - Les champs multilignes se terminent quand un nouveau label est reconnu
  - La casse des labels est ignorÃ©e (nom commun = Nom commun = NOM COMMUN)
  - Les champs inconnus sont ignorÃ©s silencieusement

Usage :
  from extract_fiches import extraire_fiche, importer_dossier
  plante = extraire_fiche("fiches/Ortie.docx")
  if plante:
      sauvegarder_plante(plante)
"""

import os
import re
from docx import Document
from models import creer_plante, Plante
from database import CHAMPS_SPECIFIQUES

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# MAPPING LABELS â†’ ATTRIBUTS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ClÃ©   = label tel qu'il apparaÃ®t dans la fiche (insensible Ã  la casse)
# Valeur = nom de l'attribut Python dans les classes

LABELS_COMMUNS = {
    "nom commun":                "nom",
    "nom":                       "nom",
    "nom scientifique":          "latin",
    "latin":                     "latin",
    "famille botanique":         "famille",
    "famille":                   "famille",
    "biologique":                "bio",
    "bio":                       "bio",
    "propriÃ©tÃ©s":                "proprietes",
    "proprietes":                "proprietes",
    "indications":               "proprietes",
    "bÃ©nÃ©fices":                 "proprietes",
    "contre-indications":        "contre",
    "contre indications":        "contre",
    "interactions mÃ©dicamenteuses": "interactions",
    "interactions":              "interactions",
    "prÃ©cautions":               "precautions",
    "precautions":               "precautions",
    "distributeur":              "distributeur",
    "fournisseur":               "distributeur",
    "prix":                      "prix",
    "quantitÃ© en stock":         "quantite",
    "quantitÃ©":                  "quantite",
    "stock":                     "quantite",
    "lieu de stockage":          "stockage",
    "stockage":                  "stockage",
    "liens":                     "liens",
    "liens ressources":          "liens",
    "ressources":                "liens",
    "notes":                     "notes",
}

LABELS_BRUTE = {
    "partie utilisÃ©e":    "partie",
    "partie":             "partie",
    "origine":            "origine",
    "provenance":         "origine",
    "mode de prÃ©paration":"mode_preparation",
    "prÃ©paration":        "mode_preparation",
    "tempÃ©rature":        "temperature",
    "temps d'infusion":   "temps_infusion",
    "infusion":           "temps_infusion",
    "posologie":          "posologie",
    "conditionnement":    "conditionnement",
}

LABELS_COMPLEMENT = {
    "partie utilisÃ©e":    "partie",
    "partie":             "partie",
    "origine":            "origine",
    "rÃ©fÃ©rence produit":  "reference",
    "rÃ©fÃ©rence":          "reference",
    "forme":              "forme",
    "dosage":             "dosage",
    "posologie":          "posologie",
    "moment de prise":    "moment_prise",
    "moment":             "moment_prise",
    "durÃ©e de cure":      "duree_cure",
    "durÃ©e":              "duree_cure",
    "conditionnement":    "conditionnement",
}

LABELS_HE = {
    "organe distillÃ©":    "organe",
    "organe":             "organe",
    "origine":            "origine",
    "mode d'obtention":   "mode_obtention",
    "obtention":          "mode_obtention",
    "chÃ©motype":          "chemotype",
    "chemotype":          "chemotype",
    "composition":        "composition",
    "voies d'utilisation":"voies",
    "voies":              "voies",
    "prÃ©cautions par voie":"precautions_voies",
    "prÃ©cautions voies":  "precautions_voies",
    "date limite":        "dlc",
    "dlc":                "dlc",
}

LABELS_JARDIN = {
    "partie":             "partie",
    "emplacement":        "emplacement",
    "exposition":         "exposition",
    "type de sol":        "type_sol",
    "sol":                "type_sol",
    "pÃ©riode de semis":   "periode_semis",
    "semis":              "periode_semis",
    "pÃ©riode de rÃ©colte": "periode_recolte",
    "rÃ©colte":            "periode_recolte",
    "vivace":             "vivace",
    "hivernage":          "hivernage",
    "entretien":          "entretien",
}

TYPE_MAP_LABELS = {
    "brute":      {**LABELS_COMMUNS, **LABELS_BRUTE},
    "complement": {**LABELS_COMMUNS, **LABELS_COMPLEMENT},
    "he":         {**LABELS_COMMUNS, **LABELS_HE},
    "jardin":     {**LABELS_COMMUNS, **LABELS_JARDIN},
}

# Synonymes pour le champ "type" dans la fiche
TYPE_SYNONYMES = {
    "plante brute":        "brute",
    "brute":               "brute",
    "tisane":              "brute",
    "complÃ©ment":          "complement",
    "complement":          "complement",
    "complÃ©ments":         "complement",
    "huile essentielle":   "he",
    "he":                  "he",
    "huile":               "he",
    "jardin":              "jardin",
    "plante jardin":       "jardin",
    "jardin potager":      "jardin",
}


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# EXTRACTION
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def _normaliser_label(texte: str) -> str:
    """Normalise un label pour la comparaison (minuscules, sans accents doublÃ©s)."""
    return texte.strip().lower().rstrip(":")


def _valeur_bool(texte: str) -> bool:
    """InterprÃ¨te 'oui/non/true/false/1/0' en boolÃ©en."""
    return texte.strip().lower() in ("oui", "yes", "true", "1", "vrai")


def extraire_fiche(fichier_path: str) -> Plante | None:
    """
    Lit un fichier .docx structurÃ© et retourne un objet Plante.

    Retourne None si :
      - le fichier n'existe pas
      - le champ 'nom commun' est absent
      - le champ 'type' est absent ou non reconnu

    Affiche des warnings pour les champs non reconnus.
    """
    if not os.path.exists(fichier_path):
        print(f"âŒ Fichier introuvable : {fichier_path}")
        return None

    try:
        doc = Document(fichier_path)
    except Exception as e:
        print(f"âŒ Impossible de lire {fichier_path} : {e}")
        return None

    # â”€â”€ PremiÃ¨re passe : dÃ©tecter le type â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    type_detecte = None
    for para in doc.paragraphs:
        texte = para.text.strip()
        if not texte or ":" not in texte:
            continue
        label, _, valeur = texte.partition(":")
        if _normaliser_label(label) == "type":
            type_detecte = TYPE_SYNONYMES.get(valeur.strip().lower())
            if type_detecte:
                break

    if not type_detecte:
        print(f"âš ï¸  Type non reconnu dans {os.path.basename(fichier_path)}")
        return None

    # â”€â”€ DeuxiÃ¨me passe : extraire tous les champs â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    labels_map = TYPE_MAP_LABELS[type_detecte]
    donnees: dict[str, str] = {}
    champ_courant: str | None = None
    valeur_courante: list[str] = []

    def _sauver_champ():
        if champ_courant:
            donnees[champ_courant] = "\n".join(valeur_courante).strip()

    for para in doc.paragraphs:
        texte = para.text.strip()
        if not texte:
            if champ_courant:
                valeur_courante.append("")
            continue

        if ":" in texte:
            label_brut, _, valeur = texte.partition(":")
            label_norm = _normaliser_label(label_brut)
            attribut = labels_map.get(label_norm)

            if attribut:
                _sauver_champ()
                champ_courant = attribut
                valeur_courante = [valeur.strip()]
                continue
            # Pas un label reconnu â†’ continuation du champ courant
        if champ_courant:
            valeur_courante.append(texte)

    _sauver_champ()  # sauver le dernier champ

    # â”€â”€ Validation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    nom = donnees.get("nom", "").strip()
    if not nom:
        print(f"âš ï¸  Champ 'Nom commun' manquant dans {os.path.basename(fichier_path)}")
        return None

    # â”€â”€ Construction de l'objet â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    obj = creer_plante(type_detecte)
    for attribut, valeur in donnees.items():
        if not hasattr(obj, attribut):
            continue
        if attribut in ("bio", "vivace"):
            setattr(obj, attribut, _valeur_bool(valeur))
        else:
            setattr(obj, attribut, valeur)

    print(f"âœ… Fiche extraite : {nom} ({type_detecte})")
    return obj


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# IMPORT EN LOT
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def importer_dossier(dossier: str) -> tuple[list[tuple[Plante, str]], list[str]]:
    """
    Parcourt le sous-dossier A_traiter/ et extrait toutes les fiches .docx.
    Les fiches extraites avec succÃ¨s sont dÃ©placÃ©es dans le dossier parent (dossier).
    Les fiches en erreur restent dans A_traiter/ pour correction.

    Retourne :
      - liste de tuples (objet Plante, chemin source) pour les succÃ¨s
      - liste des noms de fichiers en erreur
    """
    dossier_a_traiter = os.path.join(dossier, "A_traiter")

    if not os.path.isdir(dossier_a_traiter):
        print(f"âŒ Dossier introuvable : {dossier_a_traiter}")
        return [], [dossier_a_traiter]

    succes = []
    erreurs = []

    fichiers = [f for f in os.listdir(dossier_a_traiter)
                if f.lower().endswith(".docx") and not f.startswith("~")]

    if not fichiers:
        print(f"â„¹ï¸  Aucun fichier .docx trouvÃ© dans {dossier_a_traiter}")
        return [], []

    for fichier in sorted(fichiers):
        chemin = os.path.join(dossier_a_traiter, fichier)
        plante = extraire_fiche(chemin)
        if plante:
            succes.append((plante, chemin))
        else:
            erreurs.append(fichier)

    print(f"\nğŸ“Š Import terminÃ© : {len(succes)} succÃ¨s, {len(erreurs)} erreurs")
    return succes, erreurs
